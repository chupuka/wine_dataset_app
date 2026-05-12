"""
Генерация отчета в формате Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import json

# Создаем документ
doc = Document()

# Настройки страницы
sections = doc.sections
for section in sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

# Загружаем контент
with open('report_content.json', 'r', encoding='utf-8') as f:
    content = json.load(f)

# ==================== ТИТУЛЬНЫЙ ЛИСТ ====================
# Министерство
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Министерство науки и высшего образования Российской Федерации")
run.font.size = Pt(11)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("федеральное государственное бюджетное образовательное учреждение высшего образования")
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("«Российский экономический университет имени Г.В. Плеханова»")
run.font.size = Pt(10)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("МОСКОВСКИЙ ПРИБОРОСТРОИТЕЛЬНЫЙ ТЕХНИКУМ")
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Специальность: 09.02.07 Информационные системы и программирование")
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Квалификация: Программист")
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# Заголовок практической
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ПРАКТИЧЕСКАЯ РАБОТА №10")
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ПО «Технология разработки и защиты баз данных»")
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ТЕМА: Анализ текстовых данных")
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(f"Выполнил студент группы П-3-23: Бортникова А.П.")
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Проверил преподаватель: Чекан Д.В.")
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Москва 2026")
run.font.size = Pt(12)

# Разрыв страницы
doc.add_page_break()

# ==================== СОДЕРЖИМОЕ ====================
for section in content['sections']:
    # Заголовок раздела
    heading = doc.add_heading(section['heading'], level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Текст контента
    if 'content' in section:
        for para_text in section['content'].split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Таблица
    if 'table' in section:
        table = doc.add_table(rows=len(section['table']['rows']) + 1, cols=len(section['table']['headers']))
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        header_cells = table.rows[0].cells
        for i, header in enumerate(section['table']['headers']):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        # Данные таблицы
        for row_idx, row_data in enumerate(section['table']['rows']):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = cell_data
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    # Список (bullets)
    if 'bullets' in section:
        for bullet in section['bullets']:
            p = doc.add_paragraph(bullet, style='List Bullet')
    
    # Описание после таблицы
    if 'description' in section:
        for para_text in section['description'].strip().split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Изображение
    if 'image' in section:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(section['image'], width=Inches(6))
            
            if 'description' in section:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(section['description'].strip())
                run.italic = True
                run.font.size = Pt(10)
        except Exception as e:
            print(f"Ошибка при добавлении изображения {section['image']}: {e}")
            p = doc.add_paragraph(f"[Изображение: {section['image']}]")
    
    doc.add_paragraph()

# Сохраняем документ
output_path = 'Практическая10_отчет.docx'
doc.save(output_path)
print(f"✅ Отчет сохранен: {output_path}")